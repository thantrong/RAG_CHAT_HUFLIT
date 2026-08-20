"""Trích xuất text từ các loại file tải về (PDF, DOCX...).

PDF: phát hiện BẢNG và trích riêng từng hàng/cột (phân cách " | ")
để số liệu không bị dính vào nhau (vd: "10.000.0002 Trần Ngọc...").
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_file(path: Path) -> Optional[str]:
    """Trích text từ file dựa trên extension.

    Hỗ trợ: pdf, docx. Trả về None nếu không hỗ trợ hoặc lỗi.
    """
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".docx":
            return _extract_docx(path)
        if ext in (".doc", ".xls", ".xlsx", ".zip", ".rar"):
            logger.info("Chưa hỗ trợ trích text từ %s: %s", ext, path.name)
            return None

        if ext in (".txt", ".md", ".csv"):
            return path.read_text(encoding="utf-8", errors="replace")
        return None
    except Exception as e:
        logger.warning("Lỗi trích text %s: %s", path, e)
        return None


def _extract_pdf(path: Path) -> str:
    """Trích text PDF, xử lý BẢNG riêng để không bị dính cell.

    Vấn đề cũ: page.extract_text() gộp các cell trong bảng thành 1 dòng
    không phân cách, khiến số liệu dính nhau (vd: "10.000.0002 Trần Ngọc..."
    = số tiền dòng trên + STT dòng dưới dính vào nhau).

    Cách sửa:
      - Phát hiện bảng bằng page.find_tables().
      - Trích text NGOÀI bảng theo từng vùng dọc (giữ thứ tự đọc).
      - Trích BẢNG riêng: mỗi hàng là 1 dòng, các cột phân cách bằng " | ".
    """
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = _extract_pdf_page(page)
            if page_text and page_text.strip():
                parts.append(page_text)
    return "\n\n".join(parts)


def _extract_pdf_page(page) -> str:
    """Trích text 1 trang PDF, tách bảng riêng và giữ thứ tự dọc."""
    try:
        tables = page.find_tables()
    except Exception:
        tables = []


    if not tables:
        return page.extract_text() or ""


    tables_sorted = sorted(tables, key=lambda t: t.bbox[1])

    blocks: list[tuple[float, str]] = []
    page_w, page_h = page.width, page.height
    current_top = 0.0
    GAP = 2.0

    for table in tables_sorted:
        _x0, top, _x1, bottom = table.bbox


        if top > current_top + GAP:
            try:
                region = page.crop((0, current_top, page_w, top))
                text = region.extract_text()
                if text and text.strip():
                    blocks.append((current_top, text.strip()))
            except Exception:
                pass


        table_lines = _format_table(table)
        if table_lines:
            blocks.append((top, "\n".join(table_lines)))

        current_top = bottom


    if current_top < page_h - GAP:
        try:
            region = page.crop((0, current_top, page_w, page_h))
            text = region.extract_text()
            if text and text.strip():
                blocks.append((current_top, text.strip()))
        except Exception:
            pass

    return "\n".join(text for _, text in blocks)


def _format_table(table) -> list[str]:
    """Định dạng bảng: mỗi hàng là 1 dòng, các cột phân cách bằng ' | '."""
    try:
        rows = table.extract()
    except Exception:
        return []
    lines: list[str] = []
    for row in rows:
        if not row:
            continue
        cells = []
        for c in row:
            if c is None:
                c = ""

            c = " ".join(str(c).split())
            cells.append(c)
        line = " | ".join(cells).strip(" |")
        if line:
            lines.append(line)
    return lines


def _extract_docx(path: Path) -> str:
    import docx

    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cells).strip(" |")
            if line:
                parts.append(line)
    return "\n".join(parts)
